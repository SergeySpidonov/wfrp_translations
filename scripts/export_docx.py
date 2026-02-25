"""
export_docx.py — экспортирует переводы в DOCX с двумя столбцами: EN | RU.
Использование:
    python scripts/export_docx.py              # все переведённые страницы
    python scripts/export_docx.py 3 10        # страницы 3-10
"""
import sys, os, re
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TRANS_DIR   = "translations/death-on-the-reik"
OUTPUT_FILE = "Смерть на Рейке — билингвальный текст.docx"

def parse_file(filepath):
    """Возвращает список блоков [(type, text)]"""
    with open(filepath, encoding="utf-8") as f:
        content = f.read()
    blocks = []
    for raw in re.split(r"\n\s*\n", content.strip()):
        raw = raw.strip()
        if not raw or raw.startswith("=== "):
            continue
        if raw.startswith("[H1] "):
            blocks.append(("h1", raw[5:]))
        elif raw.startswith("[H2] "):
            blocks.append(("h2", raw[5:]))
        elif raw.startswith("[B] "):
            blocks.append(("bold", raw[4:]))
        elif raw.startswith("[I] "):
            blocks.append(("italic", raw[4:]))
        else:
            blocks.append(("body", raw))
    return blocks

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_text_to_cell(cell, text, block_type, doc):
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.clear()

    run = para.add_run(text)
    font = run.font

    if block_type == "h1":
        font.size = Pt(13)
        font.bold = True
        font.color.rgb = RGBColor(0x4A, 0x00, 0x00)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif block_type == "h2":
        font.size = Pt(10)
        font.bold = True
        font.color.rgb = RGBColor(0x2C, 0x2C, 0x6E)
    elif block_type == "bold":
        font.size = Pt(8.5)
        font.bold = True
    elif block_type == "italic":
        font.size = Pt(8.5)
        font.italic = True
    else:
        font.size = Pt(8.5)

    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(2)

def main():
    if len(sys.argv) == 3:
        start, end = int(sys.argv[1]), int(sys.argv[2])
    else:
        start, end = 1, 200

    doc = Document()

    # Настройка полей страницы
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Стиль документа
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(8.5)

    pages_added = 0

    for page_num in range(start, end + 1):
        en_path = os.path.join(TRANS_DIR, f"page{page_num:03d}_en.txt")
        ru_path = os.path.join(TRANS_DIR, f"page{page_num:03d}_ru.txt")

        if not os.path.exists(en_path) or not os.path.exists(ru_path):
            continue

        en_blocks = parse_file(en_path)
        ru_blocks = parse_file(ru_path)

        if not en_blocks:
            continue

        # Заголовок страницы
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(8 if pages_added > 0 else 0)
        run = heading.add_run(f"─── Страница {page_num} ───")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Таблица с двумя столбцами
        count = max(len(en_blocks), len(ru_blocks))
        table = doc.add_table(rows=count, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Ширина колонок: 50/50
        for row in table.rows:
            row.cells[0].width = Cm(13.0)
            row.cells[1].width = Cm(13.0)

        for i in range(count):
            cell_en = table.cell(i, 0)
            cell_ru = table.cell(i, 1)

            if i < len(en_blocks):
                btype, btext = en_blocks[i]
                # Лёгкий кремовый фон для EN
                set_cell_bg(cell_en, "FFFEF0")
                add_text_to_cell(cell_en, btext, btype, doc)
            else:
                set_cell_bg(cell_en, "F5F5F5")

            if i < len(ru_blocks):
                btype, btext = ru_blocks[i]
                # Лёгкий голубой фон для RU
                set_cell_bg(cell_ru, "F0F4FF")
                add_text_to_cell(cell_ru, btext, btype, doc)
            else:
                set_cell_bg(cell_ru, "F5F5F5")

        pages_added += 1
        print(f"  Страница {page_num}: {len(en_blocks)} EN / {len(ru_blocks)} RU блоков")

    if pages_added == 0:
        print("Нет переведённых страниц.")
        return

    try:
        doc.save(OUTPUT_FILE)
        print(f"\nГотово! {pages_added} страниц → {OUTPUT_FILE}")
    except PermissionError:
        import datetime
        ts = datetime.datetime.now().strftime("%H%M%S")
        fallback = OUTPUT_FILE.replace(".docx", f"_{ts}.docx")
        doc.save(fallback)
        print(f"\nФайл занят Word — сохранено как: {fallback}")
        print("Закройте Word и переименуйте файл, или откройте новый.")

if __name__ == "__main__":
    main()
